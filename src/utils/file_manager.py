from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.cv_model import CVModel

def save_cv_to_file(cv_model: CVModel, file_path: str):
    document = Document()

    # Add name
    document.add_heading(cv_model.name, 0)

    # Add contact info
    contact_info = f"{cv_model.email} | {cv_model.phone}"
    contact_paragraph = document.add_paragraph(contact_info)
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add summary
    document.add_heading('Professional Summary', level=1)
    document.add_paragraph(cv_model.summary)

    # Add experience
    document.add_heading('Experience', level=1)
    document.add_paragraph(cv_model.experience)

    # Add education
    document.add_heading('Education', level=1)
    document.add_paragraph(cv_model.education)

    document.save(file_path)