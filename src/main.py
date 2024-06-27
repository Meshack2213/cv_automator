from docx import Document
from docx.shared import Pt

def get_personal_details():
    print("Enter your personal details:")
    name = input("Name: ")
    email = input("Email: ")
    phone = input("Phone: ")
    address = input("Address: ")
    return name, email, phone, address

def get_education_details():
    print("\nEnter your education details:")
    degree = input("Degree: ")
    university = input("University: ")
    education_years = input("Years Attended: ")
    gpa = input("GPA: ")
    return degree, university, education_years, gpa

def get_work_experience():
    print("\nEnter your work experience details:")
    job_title = input("Job Title: ")
    company = input("Company: ")
    work_years = input("Years Worked: ")
    responsibilities = input("Responsibilities (separate with ';'): ")
    return job_title, company, work_years, responsibilities

def get_skills():
    print("\nEnter your skills (separate with ';'): ")
    skills_input = input("Skills: ")
    return skills_input

def add_personal_details(doc, name, email, phone, address):
    doc.add_heading("Personal Details", level=1)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")
    doc.add_paragraph(f"Address: {address}")

def add_education_details(doc, degree, university, education_years, gpa):
    doc.add_heading("Education", level=1)
    education = doc.add_paragraph()
    education.add_run(degree).bold = True
    education.add_run(f"\n{university}, {education_years}")
    education.add_run(f"\nGPA: {gpa}")

def add_work_experience(doc, job_title, company, work_years, responsibilities):
    doc.add_heading("Work Experience", level=1)
    experience = doc.add_paragraph()
    experience.add_run(job_title).bold = True
    experience.add_run(f"\n{company}, {work_years}")
    for responsibility in responsibilities.split(';'):
        experience.add_run(f"\n- {responsibility.strip()}")

def add_skills(doc, skills_input):
    doc.add_heading("Skills", level=1)
    skills = doc.add_paragraph()
    for skill in skills_input.split(';'):
        skills.add_run(f"- {skill.strip()}\n")

def create_cv():
    # Create a new Document
    doc = Document()

    # Add title
    title = doc.add_heading(level=0).add_run("Curriculum Vitae")
    title.font.size = Pt(24)
    title.bold = True

    # Get and add Personal Details
    name, email, phone, address = get_personal_details()
    add_personal_details(doc, name, email, phone, address)

    # Get and add Education Details
    degree, university, education_years, gpa = get_education_details()
    add_education_details(doc, degree, university, education_years, gpa)

    # Get and add Work Experience
    job_title, company, work_years, responsibilities = get_work_experience()
    add_work_experience(doc, job_title, company, work_years, responsibilities)

    # Get and add Skills
    skills_input = get_skills()
    add_skills(doc, skills_input)

    # Save the document
    doc.save("cv.docx")
    print("\nCV has been created successfully as 'cv.docx'.")

if __name__ == "__main__":
    create_cv()
